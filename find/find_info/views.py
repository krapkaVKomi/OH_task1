from django.shortcuts import render, redirect
from .models import *
from django.core.paginator import Paginator, EmptyPage, PageNotAnInteger
from django.contrib.auth import login, logout
from django.contrib.auth.decorators import login_required
from django.contrib.auth.forms import UserCreationForm
from django.contrib import messages
from django.http import JsonResponse
from .forms import UserLoginForm, NameForm
from chardet.universaldetector import UniversalDetector
from PyPDF2 import PdfReader
from datetime import datetime
from xlrd import open_workbook
import docx
import csv
import openpyxl


def encoding(path):
    detector = UniversalDetector()
    with open(path, 'rb') as fh:
        for line in fh:
            detector.feed(line)
            if detector.done:
                break
        detector.close()
    return detector.result['encoding']


def clear_text(line):
    new_line = str(line).strip() \
        .replace('\t', '') \
        .replace('.', '') \
        .replace(',', '') \
        .replace('!', '') \
        .replace('?', '') \
        .replace(';', '') \
        .replace('\n', '') \
        .replace('  ', ' ').split(' ')
    return new_line


@login_required
def index(request):
    try:
        checkboxs = Checkbox.objects.filter(user=request.user.id).order_by('-pk')[0]
    except IndexError:
        checkboxs = None

    if checkboxs:
        checkboxs = str(checkboxs.session)[1:-1]
        checkboxs = checkboxs.replace("'", "").split(', ')

    if request.method == "POST":
        doc = request.FILES["file"]
        doc_type = doc.name.split('.')
        doc_type = doc_type[-1]

        form_doc_name = NameForm(request.POST)
        doc_name = form_doc_name['your_name'].value()
        docs = Doc.objects.all()

        for name in docs:
            if str(name) == doc_name:
                doc_name = doc_name + ' (' + str(datetime.now()) + ')'
                break

        errors = []

        if doc_type == 'txt':
            doc_file = File.objects.create(name=doc.name, file=doc)
            doc_path = doc_file.file.path
            document = Doc.objects.create(name=doc_name, link=doc_path, word_table='Word' + str(doc_file.id))
            encod = encoding(path=doc_path)
            with open(doc_path, 'r', encoding=encod) as txt:
                line_count = 1
                for line in txt:
                    words = clear_text(line)
                    words_of_line = []
                    for word in words:
                        word = word.strip()
                        if len(word) > 3:
                            words_of_line.append(word)
                    try:
                        line_of_doc = LineOfDoc.objects.create(text=line, doc=document, line_number=line_count)
                    except:
                        errors.append(line)
                        line_count += 1
                        continue

                    line_count += 1
                    line = line.split(' ')
                    for word in line:
                        try:
                            WordOfDoc.objects.create(text=word, line=line_of_doc, doc=document)
                        except:
                            errors.append(word)
                            continue
            render(request, "main/index.html", {"doc_path": doc_path})
            return redirect('/')

        elif doc_type == 'docx':
            doc_file = File.objects.create(name=doc.name, file=doc)
            doc_path = doc_file.file.path
            document = Doc.objects.create(name=doc_name, link=doc_path)

            doc = docx.Document(doc_path)
            line_count = 1
            for i in range(len(doc.paragraphs)):
                line = doc.paragraphs[i].text
                words = clear_text(line)
                words_of_line = []
                for word in words:
                    word = word.strip()
                    if len(word) > 3:
                        words_of_line.append(word)
                try:
                    line_of_doc = LineOfDoc.objects.create(text=line, doc=document, line_number=line_count)
                except:
                    line_count += 1
                    errors.append(line)
                    continue

                line_count += 1
                line = line.split(' ')
                for word in line:
                    try:
                        WordOfDoc.objects.create(text=word, line=line_of_doc, doc=document)
                    except:
                        errors.append(word)
                        continue
            render(request, "main/index.html", {"doc_path": doc_path})
            return redirect('/')

        elif doc_type == 'csv':
            doc_file = File.objects.create(name=doc.name, file=doc)
            doc_path = doc_file.file.path
            document = Doc.objects.create(name=doc_name, link=doc_path)
            encod = encoding(path=doc_path)
            lines = []

            with open(doc_path, 'r', encoding=encod) as file:
                csvreader = csv.reader(file)
                for row in csvreader:
                    lines.append(row[0])

            line_count = 1
            for line in lines:
                line = line.replace(';', '')
                try:
                    line_of_doc = LineOfDoc.objects.create(text=line, doc=document, line_number=line_count)
                except:
                    line_count += 1
                    errors.append(line)
                    continue

                words = clear_text(line)
                for word in words:
                    try:
                        WordOfDoc.objects.create(text=word, line=line_of_doc, doc=document)
                    except:
                        errors.append(word)
                        continue
                line_count += 1

            render(request, "main/index.html", {"doc_path": doc_path})
            return redirect('/')

        elif doc_type == 'pdf':
            doc_file = File.objects.create(name=doc.name, file=doc)
            doc_path = doc_file.file.path
            document = Doc.objects.create(name=doc_name, link=doc_path)

            reader = PdfReader(doc_path)
            number_of_pages = len(reader.pages)
            for i in range(number_of_pages):
                page = reader.pages[i]
                lines_of_page = page.extract_text().split('\n')
                line_count = 1
                for line in lines_of_page:
                    try:
                        line_of_doc = LineOfDoc.objects.create(text=line, doc=document, line_number=line_count)
                    except:
                        line_count += 1
                        errors.append(line)
                        continue
                    line_count += 1
                    line = clear_text(line)
                    for word in line:
                        if len(word) > 3:
                            try:
                                WordOfDoc.objects.create(text=word, line=line_of_doc, doc=document)
                            except:
                                errors.append(word)
                                continue

            render(request, "main/index.html", {"doc_path": doc_path})
            return redirect('/')

        elif doc_type == 'xlsx':
            doc_file = File.objects.create(name=doc.name, file=doc)
            doc_path = doc_file.file.path
            document = Doc.objects.create(name=doc_name, link=doc_path)
            line_count = 1
            wb_obj = openpyxl.load_workbook(doc_path)
            sheets = wb_obj.worksheets
            for sheet_obj in sheets:
                max_col = sheet_obj.max_column
                max_row = sheet_obj.max_row
                for i in range(1, max_col + 1):
                    for j in range(1, max_row + 1):
                        cell_obj = sheet_obj.cell(row=j, column=i)
                        line = cell_obj.value
                        if line:
                            try:
                                line_of_doc = LineOfDoc.objects.create(text=line, doc=document, line_number=line_count)
                            except:
                                line_count += 1
                                errors.append(line)
                                continue

                            line = clear_text(line)
                            for word in line:
                                if len(word) > 3:
                                    try:
                                        WordOfDoc.objects.create(text=word, line=line_of_doc, doc=document)
                                    except:
                                        errors.append(word)
                                        continue
                            line_count += 1
            render(request, "main/index.html", {"doc_path": doc_path})
            return redirect('/')

        elif doc_type == 'xls' or doc_type =='XLS':
            doc_file = File.objects.create(name=doc.name, file=doc)
            doc_path = doc_file.file.path
            document = Doc.objects.create(name=doc_name, link=doc_path)
            encod = encoding(doc_path)
            wb = open_workbook(doc_path, encoding_override=encod)

            for s in wb.sheets():
                for row in range(s.nrows):
                    line = ''
                    line_count = 0
                    for col in range(s.ncols):
                        items = s.cell(row, col).value
                        item = ''.join(str(items).replace('\n', ' '))
                        line += item + ' '
                    line_count += 1
                    try:
                        line_of_doc = LineOfDoc.objects.create(text=line.strip(), doc=document, line_number=line_count)
                    except:
                        errors.append(line)
                        continue
                    line = clear_text(line)

                    for word in line:
                        if len(word) > 3:
                            try:
                                WordOfDoc.objects.create(text=word, line=line_of_doc, doc=document)
                            except:
                                errors.append(word)
                                continue
            render(request, "main/index.html", {"doc_path": doc_path})
            return redirect('/')

        else:
            print("ПОМИЛКА ", doc_type, doc_name)
            return redirect('/')
    else:
        docs = Doc.objects.all()
        get_docs = []

        for item in docs:
            select_file = request.GET.get(str(item))
            if select_file:
                get_docs.append(select_file)

        if len(docs) != 0:
            doc_arr = []

            if get_docs:
                checkboxs = get_docs
                for i in docs:
                    if i.name not in checkboxs:
                        doc_arr.append(i.name)

            else:
                if checkboxs:
                    for i in docs:
                        if i.name not in checkboxs:
                            doc_arr.append(i.name)
                else:
                    for i in docs:
                        doc_arr.append(i.name)

            if checkboxs != None and len(checkboxs) != 0:
                list_of_docs = {'docs': doc_arr, 'checkboxs': checkboxs}

            else:
                arr = []
                for i in docs:
                    arr.append(i.name)
                list_of_docs = {'docs': arr, 'checkboxs': False}
        else:
            list_of_docs = None

        query = request.GET.get('q')
        if get_docs:
            session = Checkbox.objects.create(user=request.user.id, session=get_docs)

        page = request.GET.get('page')

        if query:
            all = []
            for doc in get_docs:
                select_file_id = Doc.objects.filter(name=doc)[0:]
                select_file_id = select_file_id.values('id').get()
                select_file_id = select_file_id['id']

                #count_in_doc = (WordOfDoc.objects.filter(doc_id=select_file_id) & WordOfDoc.objects.filter(text=query)).count()

                lines_of_tible = WordOfDoc.objects.filter(doc_id=select_file_id) & WordOfDoc.objects.filter(text__istartswith=query)
                for i in lines_of_tible:
                    all.append(i)

            paginator = Paginator(all, 30)  # 30 posts per page

            try:
                posts = paginator.page(page)
            except PageNotAnInteger:
                posts = paginator.page(1)
            except EmptyPage:
                posts = paginator.page(paginator.num_pages)

            search = ''
            for i in get_docs:
                search += f"&{i}={i}"

            if page:
                if paginator.page_range.stop > 10:
                    start = f'{int(page)-1}'
                    if start != '0':
                        if paginator.page_range.stop - int(page) > 10:
                            page = f'{start}:{int(page)+10}'
                        else:
                            page = f'{paginator.page_range.stop - 10}:{paginator.page_range.stop}'

                    else:
                        page = f'{int(page)}:{int(page) + 10}'

                elif paginator.page_range.stop <= 10:
                    page = f'{int(page)-1}:{paginator.page_range.stop}'

            else:
                page = f'1:10'

            context = {
                'search': search,
                'table': True,
                'posts': posts,
                'page': page,
                'list_of_docs': list_of_docs,
                'page_num': page
            }

            return render(request, "main/index.html", context)

        else:
            context = {'table': False, 'list_of_docs': list_of_docs}
            return render(request, "main/index.html", context)


def register(request):
    if request.method == 'POST':
        form = UserCreationForm(request.POST)
        if form.is_valid():
            user = form.save()
            login(request, user)
            messages.success(request, 'Registration is successful')
            return redirect('/')
        else:
            messages.error(request, 'Registration error')
    else:
        form = UserCreationForm()
    return render(request, 'main/register.html', {"form": form})


def my_view(request):
    if request.method == 'POST':
        doc_names = []
        docs = Doc.objects.all()
        for item in docs:
            doc_names.append(item.name)

        test_input = request.POST.get('testInput')
        test_input = str(test_input)
        file_extension = request.POST.get('fileExtension')

        if not file_extension:
            data = {'error': 'upload your file!'}

        elif not test_input:
            data = {'error': 'Add name for your file!'}

        elif file_extension not in ['txt', 'pdf', 'csv', 'docx', 'xlsx', 'xls', 'XLS']:
            data = {'error': 'We cannot read the file of this extension!'}

        elif test_input in doc_names:
            data = {'error': 'This file name is already used, try another one!'}

        elif len(test_input) > 50:
            data = {'error': 'The file name is too long!'}

        else:
            num_chars = len(test_input)
            shifted_input = test_input[num_chars - 10:] + test_input[:num_chars - 10]
            data = {'message': 'We can read and download your file!', 'shiftedInput': shifted_input}

        return JsonResponse(data)


def user_login(request):
    if request.method == 'POST':
        form = UserLoginForm(data=request.POST)
        if form.is_valid():
            user = form.get_user()
            login(request, user)
            return redirect('/')
    else:
        form = UserLoginForm()
    return render(request, 'main/login.html', {"form": form})


def user_logout(request):
    logout(request)
    return redirect('/')

