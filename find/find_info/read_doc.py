from PyPDF2 import PdfReader
from chardet import UniversalDetector
from xlrd import open_workbook
import docx
import csv
from datetime import datetime
import openpyxl
import comtypes.client


def convert_doc_to_pdf(path, name):
    wdFormatPDF = 17
    out_file = path.split('\\')
    out_file[-1] = f'{name}.pdf'
    out_file = '\\'.join(out_file)
    word = comtypes.client.CreateObject('Word.Application')
    doc = word.Documents.Open(path)
    doc.SaveAs(out_file, FileFormat=wdFormatPDF)


def encoding(path):
    detector = UniversalDetector()
    with open(path, 'rb') as fh:
        for line in fh:
            detector.feed(line)
            if detector.done:
                break
        detector.close()
    return detector.result['encoding']


def clear_text(line):
    new_line = line.strip() \
        .replace('\t', '') \
        .replace('.', '') \
        .replace(',', '') \
        .replace('!', '') \
        .replace('?', '') \
        .replace(';', '') \
        .replace('\n', '') \
        .replace('  ', ' ').split(' ')
    return new_line


def read_xls(path):
    encod = encoding(path)
    wb = open_workbook(path, encoding_override=encod)

    for s in wb.sheets():
        for row in range(s.nrows):
            line = ''
            for col in range(s.ncols):
                items = s.cell(row, col).value
                item = ''.join(str(items).replace('\n', ' '))
                line += item + ' '
                print(line)
                for word in line:
                    clear_text(word)


def read_docx(path):
    test = path.split('.')
    if len(test) != 2:
        exit()

    doc = docx.Document(path)
    for i in range(len(doc.paragraphs)):
        line = doc.paragraphs[i].text
        print(line)
        words = clear_text(line)
        for word in words:
            word = word.strip()


def read_pdf(path):
    reader = PdfReader(path)

    number_of_pages = len(reader.pages)
    for i in range(number_of_pages):
        page = reader.pages[i]
        lines_of_page = page.extract_text().split('\n')
        for line in lines_of_page:
            print(line)
            line = clear_text(line)


def read_txt(path):
    encod = encoding(path=path)
    with open(path, 'r', encoding=encod) as txt:
        for line in txt:
            print(line)
            words = clear_text(line)
            for word in words:
                word = word.strip()


def read_csv(path):
    encod = encoding(path=path)
    lines = []
    with open(path, 'r', encoding=encod) as file:
        csvreader = csv.reader(file)
        for row in csvreader:
            lines.append(row[0])

    for line in lines:
        line = line.replace(';', '')
        line = clear_text(line)


def read_xlsx(path):
    wb_obj = openpyxl.load_workbook(path)
    sheets = wb_obj.worksheets
    for sheet_obj in sheets:
        max_col = sheet_obj.max_column
        max_row = sheet_obj.max_row
        for i in range(1, max_col + 1):
            for j in range(1, max_row + 1):
                cell_obj = sheet_obj.cell(row=j, column=i)
                line = cell_obj.value
                if line:
                    print(line)
                    clear_text(str(line))


start = datetime.now()
# read_xlsx(path='hard.xlsx')
convert_doc_to_pdf(r'D:\Users\User\Have_to\one\find\find_info\hard2.doc', name='namePDF2')
# read_xls('hard.xls')
# read_word('hard.doc')  # need fix
# read_word('hard.docx')
# read_pdf('hard.pdf')
# read_txt('hard.txt')
# read_csv('hard.csv')


duration = datetime.now() - start
print(f'Тривалість роботи програми: {duration}')
